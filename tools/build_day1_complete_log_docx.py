from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "InterviewPrep_AI_Day_1_Project_Log.docx"


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run(run, size=14 if level == 1 else 12, bold=True, color="2E74B5")
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(text), size=10)


def shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(2.0), Inches(4.5)]

    for idx, header in enumerate(["Area", "What We Built"]):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "F2F4F7")
        set_run(cell.paragraphs[0].add_run(header), bold=True, color="0B2545")

    for area, detail in rows:
        cells = table.add_row().cells
        for idx, value in enumerate([area, detail]):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_run(cells[idx].paragraphs[0].add_run(value), size=9)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    set_run(title.add_run("InterviewPrep AI - Day 1 Project Log"), size=20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    set_run(subtitle.add_run("Foundation milestone for the AI interview preparation platform"), size=10, color="555555")

    add_heading(doc, "Goal")
    doc.add_paragraph(
        "Build the first clean backend foundation for a resume-ready Python AI project. "
        "The app analyzes job descriptions or URLs, saves jobs, creates prep plans, generates exams, and prepares for mock interviews."
    )

    add_heading(doc, "Completed Today")
    add_table(
        doc,
        [
            ("Backend", "FastAPI app with health, job analysis, and prep-plan routes."),
            ("AI Layer", "OpenAI-ready job analyzer with local fallback and URL/page-text support."),
            ("Planning", "Day-by-day prep plan generation based on interview date, job skills, and daily study time."),
            ("Persistence", "Saved jobs, analyses, prep plans, tasks, exams, answers, and interview experiences."),
            ("Extension", "Chrome extension that saves the current job page from LinkedIn, Handshake, or similar sites."),
            ("Frontend", "React dashboard for job input, saved jobs, prep timeline, exam actions, and activity."),
            ("Mock", "Mock interview sessions with first question, answer scoring, feedback, and follow-up prompts."),
            ("Testing", "Tests for job analysis, URL extraction, prep planning, exams, persistence, and database creation."),
            ("Docs", "Clean README files, Day 1 notes, and a one-page Word project log."),
        ],
    )

    add_heading(doc, "How To Review The Code")
    for item in [
        "Start with backend/app/main.py to see how the API is assembled.",
        "Read backend/app/routers to see the available endpoints.",
        "Read backend/app/services/job_analyzer.py for the AI analysis flow.",
        "Read backend/app/services/planner.py for the schedule-generation logic.",
        "Read backend/app/services/exam_service.py for exam generation and scoring.",
        "Read backend/app/models/interview.py for the database design.",
        "Open browser-extension to review the Chrome job-saver popup.",
        "Open frontend/src/main.jsx to review the first dashboard UI.",
        "Read backend/tests to understand what behavior is currently verified.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Input Needed Later")
    for item in [
        "OpenAI API key when we want to test the real AI path.",
        "A real job posting URL to test with realistic data.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Remaining Work")
    for item in [
        "Connect the real OpenAI API for stronger job analysis, question generation, feedback, and mock interview follow-ups.",
        "Improve the frontend with saved job details, prep plan details, exam-taking screens, mock interview chat, and interview-experience forms.",
        "Polish the database layer with PostgreSQL, cleaner migrations, search/filtering, and user-specific saved progress.",
        "Add authentication so each user can save their own jobs, plans, exams, and progress.",
        "Test and polish the Chrome extension on LinkedIn, Handshake, and other job sites.",
        "Deploy the backend, frontend, and database, then create a demo video and resume-ready README.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Next Step")
    doc.add_paragraph(
        "Improve the frontend around saved plans, exams, and mock interview conversation history."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
