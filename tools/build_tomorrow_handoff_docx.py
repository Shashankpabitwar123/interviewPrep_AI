from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "docs" / "tomorrow-handoff.md"
DOCX_PATH = ROOT / "docs" / "InterviewPrep_AI_Tomorrow_Handoff.docx"


def main() -> None:
    document = Document()
    configure_document(document)

    in_code_block = False
    code_lines = []

    for raw_line in MARKDOWN_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(37, 99, 235)
            paragraph.space_after = Pt(10)
            continue

        if line.startswith("## "):
            paragraph = document.add_heading(line[3:], level=1)
            paragraph.space_before = Pt(12)
            paragraph.space_after = Pt(6)
            continue

        if line.startswith("### "):
            paragraph = document.add_heading(line[4:], level=2)
            paragraph.space_before = Pt(8)
            paragraph.space_after = Pt(4)
            continue

        if line.startswith("- "):
            document.add_paragraph(clean_inline_code(line[2:]), style="List Bullet")
            continue

        if line[:3].replace(".", "").isdigit() and ". " in line[:5]:
            _, text = line.split(". ", 1)
            document.add_paragraph(clean_inline_code(text), style="List Number")
            continue

        document.add_paragraph(clean_inline_code(line))

    note = document.add_paragraph()
    note.add_run("Latest update included: ").bold = True
    note.add_run("login/create account UI, backend user storage, password hashing, and updated remaining-work list.")

    document.save(DOCX_PATH)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(17, 24, 39)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(52, 64, 84)


def clean_inline_code(text: str) -> str:
    return text.replace("`", "")


if __name__ == "__main__":
    main()
